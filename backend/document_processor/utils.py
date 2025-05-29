import os
import PyPDF2
import docx
from io import BytesIO

def extract_text_from_pdf(file_path_or_stream):
    """
    Extract text from a PDF file.
    
    Args:
        file_path_or_stream: Path to the PDF file or a file-like object
        
    Returns:
        str: Extracted text from the PDF
    """
    text = ""
    try:
        # Check if input is a file path or a file-like object
        if isinstance(file_path_or_stream, str) and os.path.exists(file_path_or_stream):
            with open(file_path_or_stream, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                    # Add a newline between pages for better readability
                    if page_num < len(pdf_reader.pages) - 1:
                        text += "\n\n"
        elif hasattr(file_path_or_stream, 'read'):
            # It's a file-like object
            pdf_reader = PyPDF2.PdfReader(file_path_or_stream)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
                # Add a newline between pages for better readability
                if page_num < len(pdf_reader.pages) - 1:
                    text += "\n\n"
        else:
            return "Invalid PDF input"
    except Exception as e:
        text = f"Error extracting text from PDF: {str(e)}"
    
    return text

def extract_text_from_docx(file_path_or_stream):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path_or_stream: Path to the DOCX file or a file-like object
        
    Returns:
        str: Extracted text from the DOCX
    """
    text = ""
    try:
        # Check if input is a file path or a file-like object
        if isinstance(file_path_or_stream, str) and os.path.exists(file_path_or_stream):
            doc = docx.Document(file_path_or_stream)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
        elif hasattr(file_path_or_stream, 'read'):
            # It's a file-like object
            doc = docx.Document(file_path_or_stream)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
        else:
            return "Invalid DOCX input"
    except Exception as e:
        text = f"Error extracting text from DOCX: {str(e)}"
    
    return text

def extract_text_from_doc(file_path_or_stream):
    """
    Extract text from a DOC file.
    Note: This is a fallback method as direct DOC extraction is more complex.
    For production use, consider using a more robust solution like textract or an external service.
    
    Args:
        file_path_or_stream: Path to the DOC file or a file-like object
        
    Returns:
        str: Message indicating limitation of DOC extraction
    """
    return "Direct extraction from DOC files is not supported. Please convert to DOCX format for better results."

def extract_text(file_path_or_stream, file_type):
    """
    Extract text from a document based on its file type.
    
    Args:
        file_path_or_stream: Path to the document file or a file-like object
        file_type: Type of the document (pdf, docx, doc)
        
    Returns:
        str: Extracted text from the document
    """
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path_or_stream)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path_or_stream)
    elif file_type == 'doc':
        return extract_text_from_doc(file_path_or_stream)
    else:
        return "Unsupported file type. Please upload a PDF or DOCX file." 